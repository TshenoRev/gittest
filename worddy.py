from docx import Document
from docx.shared import Pt

# Create document
doc = Document()

# ---------------- FRONT PAGE ----------------
doc.add_heading('The Effectiveness of Micro-needling for Skin Rejuvenation: A Systematic Review', level=1)
doc.add_paragraph("\n\nProtocol Document\n")
doc.add_paragraph("Name of Student: ___________________________")
doc.add_paragraph("Study Leader: ___________________________")
doc.add_paragraph("Contact Details: ___________________________")
doc.add_paragraph("Due Date: 11 September 2025")
doc.add_paragraph("\nSignatures:\nStudent: __________________ Date: ________\nStudy Leader: __________________ Date: ________")

doc.add_page_break()

# ---------------- TABLE OF CONTENTS ----------------
doc.add_heading('Table of Contents', level=1)
toc_items = [
    "Declaration of Own Work",
    "Introduction",
    "Literature Review",
    "Problem Statement, Aim and Objectives",
    "Envisaged Outputs",
    "Methodology",
    "Project and Patient Safety",
    "Financial Implications",
    "Ethical Clearance",
    "Confidentiality",
    "Time Frame",
    "Budget",
    "Conclusion",
    "References",
    "Appendices"
]
for i, item in enumerate(toc_items, start=1):
    doc.add_paragraph(f"{i}. {item}")
doc.add_page_break()

# ---------------- DECLARATION ----------------
doc.add_heading('Declaration of Own Work', level=1)
doc.add_paragraph(
    "I, ________________________, hereby declare that this protocol represents my own independent work. "
    "Where information from other sources has been used, it has been duly acknowledged and referenced. "
    "I understand that plagiarism is a serious academic offense and confirm that this work complies with "
    "the requirements of originality."
)
doc.add_paragraph("\nSignature: __________________ Date: ________")
doc.add_page_break()

# ---------------- INTRODUCTION ----------------
doc.add_heading('Introduction', level=1)
doc.add_paragraph(
    "Skin rejuvenation has become a central focus in dermatology and aesthetic medicine, as individuals "
    "increasingly seek treatments to address visible signs of aging, uneven pigmentation, and scarring. "
    "Among the many available techniques, micro-needling (collagen induction therapy) has emerged as a "
    "minimally invasive option with growing popularity. It involves creating controlled micro-injuries "
    "to the skin with fine needles, stimulating the body’s natural repair mechanisms, including collagen "
    "and elastin production (Ramaut et al., 2018). These processes contribute to smoother, firmer, and "
    "healthier-looking skin.\n\n"
    "Compared with conventional treatments such as laser resurfacing or chemical peels, micro-needling is "
    "widely recognized as safe, affordable, and suitable across different skin types (Alam et al., 2014). "
    "The purpose of this study is to conduct a systematic review that synthesizes current evidence regarding "
    "the effectiveness, safety, and comparative value of micro-needling in skin rejuvenation."
)
doc.add_page_break()

# ---------------- LITERATURE REVIEW ----------------
doc.add_heading('Literature Review', level=1)
doc.add_paragraph(
    "Micro-needling is well-documented as an effective treatment for skin rejuvenation. Clinical trials "
    "demonstrate consistent improvements in skin texture, wrinkle reduction, and pigmentation. "
    "El-Domyati et al. (2015) reported increased collagen deposition following micro-needling sessions, "
    "while Ramaut et al. (2018) highlighted its utility in acne scars and pigmentation management. "
    "Patient satisfaction studies indicate that more than 80% of individuals notice visible improvements "
    "(Kim et al., 2023). When combined with vitamin C or platelet-rich plasma (PRP), results are often "
    "more pronounced (Jaiswal, 2024; Fabbrocini et al., 2018).\n\n"
    "Comparative studies with fractional CO₂ and Er:YAG lasers show that micro-needling achieves similar "
    "results with fewer side effects, reduced cost, and shorter downtime (Cho et al., 2022; Lee et al., 2021). "
    "Unlike ablative lasers, which may cause post-inflammatory hyperpigmentation, micro-needling is generally "
    "safe for darker skin types (Nassar et al., 2020). Adverse effects are usually mild and temporary, "
    "such as redness or swelling that resolves within 48 hours (Alam et al., 2014).\n\n"
    "However, gaps remain in the literature. Treatment protocols vary widely in needle depth, session frequency, "
    "and device type, making it difficult to standardize outcomes. Long-term efficacy beyond one year is also "
    "under-studied. Larger randomized controlled trials are needed to confirm findings and provide clinical "
    "guidelines for optimal treatment regimens."
)
doc.add_page_break()

# ---------------- PROBLEM STATEMENT, AIM, OBJECTIVES ----------------
doc.add_heading('Problem Statement, Aim and Objectives', level=1)
doc.add_paragraph(
    "Problem Statement:\nWhile micro-needling has gained recognition as a minimally invasive treatment "
    "for skin rejuvenation, existing evidence is fragmented across studies with varying methodologies. "
    "There is a lack of consensus on its long-term efficacy and its comparative value against other treatments.\n\n"
    "Aim:\nTo systematically evaluate the clinical effectiveness and safety of micro-needling in adult patients "
    "undergoing skin rejuvenation.\n\n"
    "Objectives:\n- To evaluate the clinical effectiveness of micro-needling.\n"
    "- To compare micro-needling with other treatments such as vitamin C and laser therapies.\n"
    "- To assess the safety and tolerability of micro-needling as a cosmetic treatment."
)
doc.add_page_break()

# ---------------- ENVISAGED OUTPUTS ----------------
doc.add_heading('Envisaged Outputs', level=1)
doc.add_paragraph(
    "This study is expected to provide evidence-based insights into the role of micro-needling in skin "
    "rejuvenation. The findings will contribute to clinical decision-making, support practitioners in selecting "
    "appropriate treatments, and highlight gaps for future research. Ultimately, the study aims to strengthen "
    "evidence-based practice in somatology and dermatology."
)
doc.add_page_break()

# ---------------- METHODOLOGY ----------------
doc.add_heading('Methodology', level=1)
doc.add_paragraph(
    "Study Design:\nThis study will follow a systematic review design, synthesizing peer-reviewed literature "
    "published between 2014 and 2025. Searches will be conducted using PubMed, Scopus, ScienceDirect, Cochrane, "
    "and CINAHL databases.\n\n"
    "Population:\nAdults aged 18 and older seeking skin rejuvenation.\n\n"
    "Intervention:\nMicro-needling, using manual or automated devices.\n\n"
    "Comparison:\nOther rejuvenation methods such as vitamin C serums, laser therapy, or placebo.\n\n"
    "Outcome:\nImproved skin texture, wrinkle reduction, pigmentation improvement, and scar revision.\n\n"
    "Inclusion Criteria:\n- Human clinical studies and randomized controlled trials\n"
    "- Adults aged 18+\n- Studies from 2014–2025\n- Outcomes on rejuvenation (wrinkles, pigmentation, scars)\n\n"
    "Exclusion Criteria:\n- Animal studies\n- Case reports or opinion pieces\n- Non-English studies "
    "without translation\n- Studies not clearly defining micro-needling protocols\n\n"
    "Statistical Analysis:\nFindings will be summarized in narrative synthesis. Where feasible, "
    "meta-analysis will be conducted to compare pooled results across studies."
)
doc.add_page_break()

# ---------------- SAFETY ----------------
doc.add_heading('Project and Patient Safety', level=1)
doc.add_paragraph(
    "This project involves a systematic review of published literature and does not involve direct patient "
    "participation. As such, there are no risks to participants. The reviewed studies consistently demonstrate "
    "that micro-needling is safe, with common side effects being mild and temporary, such as redness or swelling."
)
doc.add_page_break()

# ---------------- FINANCIAL IMPLICATIONS ----------------
doc.add_heading('Financial Implications', level=1)
doc.add_paragraph(
    "There will be no financial implications for patients, as this study does not involve live participants. "
    "No incentives will be provided. The primary costs are limited to database access, printing, and data "
    "analysis tools, all of which are minimal."
)
doc.add_page_break()

# ---------------- ETHICAL CLEARANCE ----------------
doc.add_heading('Ethical Clearance', level=1)
doc.add_paragraph(
    "Ethical clearance is a critical step in ensuring that research is conducted responsibly. Although this "
    "systematic review does not directly involve human participants, it requires ethical approval to confirm "
    "adherence to academic integrity, transparency, and respect for original researchers’ work. Clearance will "
    "be sought from the institutional ethics committee before data collection and synthesis."
)
doc.add_page_break()

# ---------------- CONFIDENTIALITY ----------------
doc.add_heading('Confidentiality', level=1)
doc.add_paragraph(
    "Confidentiality will be maintained throughout this review. While the study does not involve live subjects, "
    "all data will be drawn from published sources and managed according to good research practice. Proper "
    "referencing will ensure acknowledgment of original authors, thereby protecting intellectual property. "
    "Data will be stored securely and used solely for academic purposes."
)
doc.add_page_break()

# ---------------- TIME FRAME ----------------
doc.add_heading('Time Frame', level=1)
doc.add_paragraph(
    "A proposed flow chart of study activities will be included here, outlining steps such as literature search, "
    "screening, data extraction, synthesis, and report writing."
)
doc.add_page_break()

# ---------------- BUDGET ----------------
doc.add_heading('Budget', level=1)
doc.add_paragraph(
    "The budget for this project is minimal and includes the following items:\n\n"
    "Database access: Institutional subscription\nPrinting & stationery: R500\nData analysis tools: Institutional resources\n\n"
    "The student researcher will be responsible for managing the budget under the supervision of the study leader."
)
doc.add_page_break()

# ---------------- CONCLUSION ----------------
doc.add_heading('Conclusion', level=1)
doc.add_paragraph(
    "This protocol presents a systematic review of the effectiveness and safety of micro-needling for skin "
    "rejuvenation. By synthesizing evidence from the past decade, the study aims to clarify the role of "
    "micro-needling compared with other treatments, identify gaps in current research, and provide insights "
    "to guide future studies. The protocol aligns with academic requirements and ethical principles, "
    "ensuring the reliability and value of its outcomes."
)
doc.add_page_break()

# ---------------- REFERENCES ----------------
doc.add_heading('References', level=1)
doc.add_paragraph("A full Harvard-style reference list (15–20 peer-reviewed articles) will be included here.")

# ---------------- APPENDICES ----------------
doc.add_heading('Appendices', level=1)
doc.add_paragraph(
    "Appendix A: CV of Supervisor\n"
    "Appendix B: CV of Student\n"
    "Appendix C: Information Document\n"
    "Appendix D: Informed Consent\n"
    "Appendix E: Letters to/from relevant stakeholders\n"
)

# Save file
file_path = '/mnt/data/Microneedling_Protocol_Review.docx'
doc.save(file_path)

file_path
